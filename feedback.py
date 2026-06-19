import streamlit as st
import google.generativeai as genai
import os
import io
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ==========================================
# 1. Page & API Configuration
# ==========================================
st.set_page_config(
    page_title="AI Tutor Feedback System",
    layout="centered",
    page_icon="🎓"
)

try:
    os.environ["GOOGLE_API_VERSION"] = "v1"
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
        genai.configure(api_key=api_key)
    else:
        st.error("❌ API Key not found in secrets.toml")
except Exception as e:
    st.error(f"❌ Configuration Error: {e}")

# ==========================================
# 2. Dynamic Model Loading
# ==========================================
@st.cache_resource
def get_available_model():
    try:
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        priority_list = ['models/gemini-1.5-flash', 'models/gemini-1.5-pro', 'gemini-1.5-flash', 'gemini-1.5-pro']
        for target in priority_list:
            if target in available_models:
                return genai.GenerativeModel(target)
        if available_models:
            return genai.GenerativeModel(available_models[0])
    except Exception as e:
        st.error(f"Model detection failed: {e}")
    return None

model = get_available_model()

# ==========================================
# 3. Word Document Generation
# ==========================================
NAVY   = RGBColor(0x1C, 0x2B, 0x4B)
ORANGE = RGBColor(0xC8, 0x5A, 0x1A)
GREY   = RGBColor(0x99, 0x99, 0x99)

# Canonical section names — the code will auto-number these 1–6
SECTIONS = [
    "Executive Summary",
    "Detailed Marking",
    "Key Gaps",
    "Next Steps",
    "Foundation (To Consolidate)",
    "Extension (To Challenge)",
]

# Lines containing any of these patterns are silently dropped
REMOVE_PATTERNS = [
    r"STARS_EDU",
    r"Academic Reporting System",
    r"Senior International Tutor",
    r"^Study Report",
    r"^Student\s*:",
    r"^Tutor\s*:",
    r"^Date\s*:",
    r"^To\s*:",
    r"^From\s*:",
    r"^Homework Feedback$",
]

def should_remove(line: str) -> bool:
    for pat in REMOVE_PATTERNS:
        if re.search(pat, line.strip(), re.IGNORECASE):
            return True
    return False

def match_section(line: str):
    """Return (section_index, canonical_label) if line is a section header, else None."""
    clean = re.sub(r'^\d+[\.\)]\s*', '', line.strip()).rstrip(':').strip().lower()
    for i, name in enumerate(SECTIONS):
        if clean == name.lower():
            return i, name
        if clean.startswith(name.lower()):
            return i, name
    return None

def is_numbered_item(line: str) -> bool:
    return bool(re.match(r'^\d+[\)\.]\s+\S', line.strip()))

def create_docx(text: str, student_name: str, teacher_name: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # --- Branding header: logo left, brand text right ---
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = hdr_table.cell(0, 0)
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
    if os.path.exists(logo_path):
        lp.add_run().add_picture(logo_path, height=Cm(1.2))
    else:
        lr = lp.add_run("PLACE LOGO")
        lr.font.size = Pt(9)
        lr.font.color.rgb = ORANGE

    right_cell = hdr_table.cell(0, 1)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("StarsEdu Online Tuition")
    rr.bold = True
    rr.font.size = Pt(12)
    rr.font.color.rgb = ORANGE

    doc.add_paragraph()

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    tr = title_para.add_run("Homework Feedback")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = NAVY

    rule_para = doc.add_paragraph()
    rule_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_para.paragraph_format.space_after = Pt(14)
    rr2 = rule_para.add_run("─────")
    rr2.font.color.rgb = ORANGE
    rr2.font.size = Pt(14)

    # --- Body ---
    # Split at "Best regards" so we control the sign-off
    body_text, _, _ = text.partition("Best regards")

    lines = body_text.split("\n")
    prev_blank = False

    for raw in lines:
        line = raw.strip()

        if should_remove(line):
            continue

        # Collapse consecutive blank lines
        if not line:
            if prev_blank:
                continue
            prev_blank = True
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        prev_blank = False

        # Section header
        m = match_section(line)
        if m is not None:
            idx, name = m
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(f"{idx + 1}.  {name}:")
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Numbered sub-item e.g. "1). Grammar…"
        if is_numbered_item(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            num_match = re.match(r'^(\d+[\)\.])(.+)', line)
            if num_match:
                r1 = p.add_run(num_match.group(1) + "  ")
                r1.font.size = Pt(11)
                r2 = p.add_run(num_match.group(2).strip())
                r2.font.size = Pt(11)
            else:
                p.add_run(line).font.size = Pt(11)
            continue

        # Bullet line
        if re.match(r'^[•\-\*]\s+', line):
            body = re.sub(r'^[•\-\*]\s+', '', line)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(body).font.size = Pt(11)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Normal paragraph
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # --- Controlled sign-off ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Best regards,").font.size = Pt(11)

    p2 = doc.add_paragraph()
    r = p2.add_run(teacher_name)
    r.bold = True
    r.font.size = Pt(11)
    p2.paragraph_format.space_before = Pt(6)

    # --- Footer ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    pPr = footer_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_el = OxmlElement("w:top")
    top_el.set(qn("w:val"), "single")
    top_el.set(qn("w:sz"), "4")
    top_el.set(qn("w:color"), "CCCCCC")
    pBdr.append(top_el)
    pPr.append(pBdr)
    fr = footer_para.add_run("STARS_EDU // Academic Reporting System")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ==========================================
# 4. Payload helper (PDF/image/Word)
# ==========================================
def build_payload_part(uploaded_file, label):
    filename = uploaded_file.name.lower()
    if filename.endswith(".docx") or filename.endswith(".doc"):
        doc = Document(io.BytesIO(uploaded_file.getvalue()))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
        return f"\n[{label}]\n" + "\n".join(lines)
    else:
        return {"mime_type": uploaded_file.type, "data": uploaded_file.getvalue()}


# ==========================================
# 5. Main Application Interface
# ==========================================
st.title("🎓 Professional Feedback Portal")
st.write("Generate high-quality feedback reports for your students instantly.")

with st.sidebar:
    st.header("📋 Report Details")
    teacher_name = st.text_input("Teacher Name")
    student_name = st.text_input("Student Name")

    st.write("---")
    st.header("📂 Upload Work")
    homework_file    = st.file_uploader(
        "Upload Student Work (PDF / Image / Word)",
        type=["pdf", "png", "jpg", "jpeg", "docx", "doc"]
    )
    mark_scheme_file = st.file_uploader(
        "Upload Mark Scheme (Optional)",
        type=["pdf", "png", "jpg", "jpeg", "docx", "doc"]
    )

if homework_file:
    st.info(f"File '{homework_file.name}' ready for analysis.")

    if st.button("🚀 Process & Generate Feedback Report", use_container_width=True):
        if not teacher_name or not student_name:
            st.warning("⚠️ Please enter both Teacher and Student names in the sidebar.")
        elif not model:
            st.error("AI model is not responding. Check your API key.")
        else:
            with st.spinner(f"Analysing {student_name}'s work…"):

                prompt = f"""You are a tutor named {teacher_name}.
Write a personalised homework feedback report addressed to your student, {student_name}.

Tone: professional, encouraging, and direct (use "I" and "you").
Today's date: {date.today().strftime("%B %d, %Y")}.

IMPORTANT — begin the report immediately with the greeting below.
Do NOT add any title block, header, or metadata before the greeting.
Do NOT write lines such as "Study Report:", "Student:", "Tutor:", "Date:", or "Homework Feedback" at the top.

Start with:

Dear {student_name},

[One warm opening sentence.]

Then include each of these sections in order, using EXACTLY these headings (no numbering — the formatting system will add numbers):

Executive Summary:
[2–3 sentences summarising overall performance.]

Detailed Marking:
[For each question or paragraph, give it a clear label and mark it [Correct], [Correct – Minor Error], or [Incorrect]. Explain what was right or wrong and give a specific suggestion.]

Key Gaps:
[Number each gap 1). 2). 3). Focus on the most important areas needing improvement.]

Next Steps:
[One brief sentence introducing the steps below.]

Foundation (To Consolidate):
[Number each step 1). 2). 3). Practical actions to reinforce core understanding.]

Extension (To Challenge):
[Number each step 1). 2). 3). Stretch tasks for deeper learning.]

[One closing encouragement sentence addressed to {student_name} by name.]

Best regards,
{teacher_name}

Rules:
- Do NOT include any title, header block, or metadata at the top.
- Do NOT write your job title, "Senior International Tutor", "STARS_EDU", or any footer.
- Do NOT add more than one blank line between sections.
- Use plain numbered items 1). 2). 3). — no bullet symbols unless listing marking points.
- Write in English only.
"""
                payload = [prompt, build_payload_part(homework_file, "Student Work")]
                if mark_scheme_file:
                    payload.append(build_payload_part(mark_scheme_file, "Mark Scheme"))

                try:
                    response = model.generate_content(payload)
                    report_text = response.text

                    docx_bytes = create_docx(report_text, student_name, teacher_name)

                    st.success(f"🎉 Feedback report for {student_name} is ready!")

                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="📥 Download Word (.docx)",
                            data=docx_bytes,
                            file_name=f"{student_name}_Feedback_{date.today()}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

                    with st.expander("Preview report text"):
                        st.text(report_text)

                except Exception as e:
                    st.error(f"Error during report generation: {e}")
else:
    st.info("Upload a student's homework to begin.")

st.markdown("---")
st.caption(f"System Date: {date.today().strftime('%Y-%m-%d')} | Powered by Gemini 1.5 Pro")
